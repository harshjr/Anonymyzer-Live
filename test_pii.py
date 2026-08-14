#!/usr/bin/env python3
"""
Test Suite for PII Document Anonymizer

Tests hybrid detection, synthetic consistency, DOCX tables, run-splitting,
hyperlinks/field codes, PDF redaction, and validation.
"""

import sys
from pathlib import Path
import tempfile
import unittest
import docx
import fitz

from pii_anonymizer import (
    PIIEngine,
    SyntheticMapper,
    process_txt,
    process_docx,
    process_pdf,
    validate,
    is_luhn_valid,
    is_valid_ipv4,
    is_valid_phone,
)


class TestPIIAnonymizer(unittest.TestCase):

    def setUp(self):
        self.engine = PIIEngine(min_score=0.55)
        self.mapper = SyntheticMapper(seed=42)

    def test_01_email_detection_and_redaction(self):
        text = "Please reach out to rashi.patil@gmail.com for any queries."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("rashi.patil@gmail.com", redacted)
        self.assertTrue(any(a["type"] == "EMAIL" for a in audit))
        self.assertIn("@", redacted)  # Has synthetic email

    def test_02_phone_detection_with_validation(self):
        text = "Contact support at +91 9876543210 or 020-67295100."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("+91 9876543210", redacted)
        self.assertTrue(any(a["type"] == "PHONE" for a in audit))

    def test_03_person_detection_and_allcaps(self):
        text = "PROMOTERS: KUSHAL SUBBAYYA HEGDE, PUSHPA KUSHAL HEGDE"
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("KUSHAL SUBBAYYA HEGDE", redacted)
        self.assertNotIn("PUSHPA KUSHAL HEGDE", redacted)
        self.assertTrue(any(a["type"] == "PERSON" for a in audit))

    def test_04_company_detection(self):
        text = "The statutory audit was conducted by Kirtane & Pandit LLP, Chartered Accountants."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("Kirtane & Pandit LLP", redacted)

    def test_05_address_detection(self):
        text = "Our office is at 201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune - 411045."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("Montreal Business Centre", redacted)

    def test_06_ssn_detection(self):
        text = "Social security number on file is 123-45-6789."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("123-45-6789", redacted)
        self.assertTrue(any(a["type"] == "SSN" for a in audit))

    def test_07_credit_card_luhn_validation(self):
        # Valid Luhn card
        valid_cc = "4532015112830366"  # Standard Visa format
        self.assertTrue(is_luhn_valid(valid_cc))
        text = f"Payment charged to card {valid_cc}."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn(valid_cc, redacted)
        self.assertTrue(any(a["type"] == "CREDIT_CARD" for a in audit))

    def test_08_dob_contextual_detection(self):
        text = "Employee John Smith, DOB: 1985-05-20, was hired yesterday."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("1985-05-20", redacted)
        self.assertTrue(any(a["type"] == "DATE_OF_BIRTH" for a in audit))

    def test_09_ip_address_validation(self):
        self.assertTrue(is_valid_ipv4("192.168.1.100"))
        self.assertFalse(is_valid_ipv4("999.999.999.999"))
        text = "Server connection from 192.168.1.100 was logged."
        redacted, audit = self.engine.redact_text(text, self.mapper)
        self.assertNotIn("192.168.1.100", redacted)
        self.assertTrue(any(a["type"] == "IP_ADDRESS" for a in audit))

    def test_10_consistent_synthetic_replacement(self):
        text = (
            "First mention: Rashi Patil and rashi.patil@gmail.com. "
            "Second mention: Rashi Patil and rashi.patil@gmail.com."
        )
        mapper = SyntheticMapper(seed=100)
        redacted, _ = self.engine.redact_text(text, mapper)
        
        # Verify both occurrences were replaced identically
        fake_name = mapper.get("PERSON", "Rashi Patil")
        fake_email = mapper.get("EMAIL", "rashi.patil@gmail.com")
        self.assertEqual(redacted.count(fake_name), 2)
        self.assertEqual(redacted.count(fake_email), 2)

    def test_11_docx_table_anonymization(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "test_table.docx"
            dst = Path(tmpdir) / "test_table_redacted.docx"

            doc = docx.Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Contact Name"
            table.cell(0, 1).text = "Sarthak Malvadkar"
            table.cell(1, 0).text = "Email Address"
            table.cell(1, 1).text = "sarthak@example.com"
            doc.save(src)

            audit = process_docx(src, dst, self.engine, self.mapper)
            self.assertTrue(len(audit) >= 2)

            # Verify by opening redacted docx
            redacted_doc = docx.Document(dst)
            cell_texts = [cell.text for row in redacted_doc.tables[0].rows for cell in row.cells]
            all_text = " ".join(cell_texts)
            self.assertNotIn("Sarthak Malvadkar", all_text)
            self.assertNotIn("sarthak@example.com", all_text)

    def test_12_docx_split_runs_anonymization(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "test_split.docx"
            dst = Path(tmpdir) / "test_split_redacted.docx"

            doc = docx.Document()
            p = doc.add_paragraph()
            r1 = p.add_run("rashi.patil")
            r2 = p.add_run("@gmail.com")
            doc.save(src)

            process_docx(src, dst, self.engine, self.mapper)

            redacted_doc = docx.Document(dst)
            full_text = redacted_doc.paragraphs[0].text
            self.assertNotIn("rashi.patil", full_text)
            self.assertNotIn("@gmail.com", full_text)
            self.assertIn("@", full_text)  # Replaced with fake email

    def test_13_pdf_redaction_and_scanned_check(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "test.pdf"
            dst = Path(tmpdir) / "test_redacted.pdf"

            # Create a simple PDF
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 72), "User: Alice Walker\nEmail: alice.walker@domain.com\nPhone: +1 555-019-2834")
            doc.save(src)
            doc.close()

            audit = process_pdf(src, dst, self.engine, self.mapper)
            self.assertTrue(len(audit) > 0)

            # Verify redacted PDF
            redacted_doc = fitz.open(dst)
            text = " ".join(page.get_text() for page in redacted_doc)
            redacted_doc.close()
            self.assertNotIn("alice.walker@domain.com", text)


    def test_14_docx_hyperlink_instrtext_anonymization(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "test_hyperlink.docx"
            dst = Path(tmpdir) / "test_hyperlink_redacted.docx"

            # Create docx with instrText XML node
            doc = docx.Document()
            p = doc.add_paragraph("Email link: ")
            doc.save(src)

            # Inject w:instrText into docx
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(src.with_name("tmp.docx"), "w") as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        data_str = data.decode("utf-8")
                        data_str = data_str.replace(
                            "Email link: ",
                            'Email link: <w:r><w:instrText xml:space="preserve"> HYPERLINK "mailto:secret.user@company.com" </w:instrText></w:r>'
                        )
                        data = data_str.encode("utf-8")
                    zout.writestr(item, data)
            src.with_name("tmp.docx").replace(src)

            audit = process_docx(src, dst, self.engine, self.mapper)
            self.assertTrue(any(a["type"] == "EMAIL" for a in audit))

            # Verify that secret.user@company.com does not exist in any part of dst
            with zipfile.ZipFile(dst, "r") as z:
                for name in z.namelist():
                    content = z.read(name).decode("utf-8", errors="ignore")
                    self.assertNotIn("secret.user@company.com", content)

    def test_15_scanned_pdf_warning(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = Path(tmpdir) / "scanned.pdf"
            dst = Path(tmpdir) / "scanned_redacted.pdf"

            # Create an empty / image-only style PDF with no text
            doc = fitz.open()
            doc.new_page()  # Blank page
            doc.save(src)
            doc.close()

            audit = process_pdf(src, dst, self.engine, self.mapper)
            self.assertEqual(len(audit), 0)


if __name__ == "__main__":
    unittest.main()
